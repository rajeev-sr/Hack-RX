import fitz  # PyMuPDF
import pdfplumber
import docx
import mailparser
from bs4 import BeautifulSoup
import pandas as pd
from PIL import Image
import pytesseract
import io
import requests

# --- Download File ---
def download_file(url):
    response = requests.get(url)
    if response.status_code == 200:
        return io.BytesIO(response.content)
    else:
        raise Exception(f"Failed to download file from {url}")

# --- Extract Text & Tables from PDF ---
def extract_from_pdf(file_bytesio):
    pdf_file=file_bytesio
    extracted_text = ""
    tables_data = []

    with pdfplumber.open(pdf_file) as pdf:
        for page_number, page in enumerate(pdf.pages):
            page_text = ""

            # Extract Tables and get their bounding boxes
            table_bboxes = []
            tables = page.extract_tables()
            for table in tables:
                # Get bounding boxes for each table
                table_settings = {"vertical_strategy": "lines", "horizontal_strategy": "lines"}
                table_bbox = page.find_tables(table_settings=table_settings)
                if table_bbox:
                    table_bboxes.append(table_bbox[0].bbox)
                tables_data.append(table)

            # Extract text but skip content inside table bounding boxes
            words = page.extract_words()
            for word in words:
                word_bbox = (word['x0'], word['top'], word['x1'], word['bottom'])
                # Check if word is inside any table bbox
                in_table = False
                for bbox in table_bboxes:
                    if (bbox[0] <= word_bbox[0] <= bbox[2]) and (bbox[1] <= word_bbox[1] <= bbox[3]):
                        in_table = True
                        break
                if not in_table:
                    page_text += word['text'] + ' '

            if page_text.strip():
                extracted_text += page_text.strip()

    return {
        "text": extracted_text.strip(),
        "tables": tables_data
    }

# --- Extract Images & OCR Summary from PDF ---
def extract_images_from_pdf(file_stream):
    doc = fitz.open(stream=file_stream, filetype="pdf")
    image_summaries = []

    for page_number in range(len(doc)):
        page = doc.load_page(page_number)
        images = page.get_images(full=True)
        
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))

            # Run OCR on image
            ocr_text = pytesseract.image_to_string(image)

            image_summaries.append({
                "page": page_number + 1,
                "image_index": img_index + 1,
                "ocr_summary": ocr_text.strip()
            })

    return image_summaries

# --- Extract Text & Tables from DOCX ---


def extract_from_docx(file_stream):
    document = docx.Document(file_stream)
    extracted_text = ""
    tables_data = []

    for para in document.paragraphs:
        if para.text.strip():
            extracted_text += para.text.strip() + "\n"

    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)

    return {
        "type": "DOCX",
        "text": extracted_text.strip(),
        "tables": tables_data
    }

# --- Extract Images & OCR Summary from DOCX ---
def extract_images_from_docx(file_stream):
    doc = docx.Document(file_stream)
    image_summaries = []
    rels = doc.part.rels

    img_index = 1
    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.target_ref:
            image_data = rel_obj.target_part.blob
            image = Image.open(io.BytesIO(image_data))

            # Run OCR on image
            ocr_text = pytesseract.image_to_string(image)

            image_summaries.append({
                "image_index": img_index,
                "ocr_summary": ocr_text.strip()
            })
            img_index += 1

    return image_summaries

# --- Extract Text & Tables from EML ---
def extract_from_eml(file_stream):
    mail = mailparser.parse_from_bytes(file_stream.read())
    extracted_text = mail.body or ""
    tables_data = []
    if mail.body_html:
        tables_data = extract_tables_from_html(mail.body_html)
    return {
        "type": "EML",
        "text": extracted_text.strip(),
        "tables": tables_data
    }

# --- Extract Tables from HTML Content (EML) ---
def extract_tables_from_html(html_content):
    tables_data = []
    soup = BeautifulSoup(html_content, 'html.parser')
    tables = soup.find_all('table')
    for table in tables:
        rows = []
        for tr in table.find_all('tr'):
            cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
            rows.append(cells)
        tables_data.append(rows)
    return tables_data

# --- Process File Based on Extension ---

def fix_google_docs_link(url):
    if 'docs.google.com' in url and '/document/' in url:
        doc_id = url.split('/d/')[1].split('/')[0]
        return f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
    return url
def process_file(url):
    

    if 'pdf' in url.lower():
        file_stream = download_file(url)
        pdf_data = extract_from_pdf(file_stream)
        file_stream.seek(0)
        image_summaries = extract_images_from_pdf(file_stream)
        pdf_data['image_summaries'] = image_summaries
        return pdf_data

    elif 'doc' in url.lower():
        url = fix_google_docs_link(url)
        file_stream = download_file(url)
        docx_data = extract_from_docx(file_stream)
        file_stream.seek(0)
        image_summaries = extract_images_from_docx(file_stream)
        docx_data['image_summaries'] = image_summaries
        return docx_data

    elif 'eml' in url.lower():
        file_stream = download_file(url)
        return extract_from_eml(file_stream)

    else:
        raise Exception("Unsupported file format. Only PDF, DOCX, and EML are supported.")
    
# print("Data preprocessing module loaded successfully.",process_file("https://hackrx.blob.core.windows.net/assets/policy.pdf?sv=2023-01-03&st=2025-07-04T09%3A11%3A24Z&se=2027-07-05T09%3A11%3A00Z&sr=b&sp=r&sig=N4a9OU0w0QXO6AOIBiu4bpl7AXvEZogeT%2FjUHNO7HzQ%3D"))
from langchain_core.documents import Document

def load_document(url,source: str = "user_data") -> list[Document]:
    
    data_list = [process_file(url)
]

    documents = []

    for data in data_list:
        doc_parts = []

        if 'text' in data and data['text']:
            doc_parts.append(data['text'])

        if 'tables' in data and data['tables']:
            for table_idx, table in enumerate(data['tables'], 1):
                table_str = ""
                for row in table:
                    row_items = [str(item) if item is not None else "" for item in row]
                    table_str += "| " + " | ".join(row_items) + " |\n"
                doc_parts.append(table_str.strip())

        if 'image_summaries' in data and data['image_summaries']:
            doc_parts.append("## Image Summaries")
            for img_summary in data['image_summaries']:
                page = img_summary.get('page', 'N/A')
                img_idx = img_summary.get('image_index', 'N/A')
                summary = img_summary.get('ocr_summary', '')
                doc_parts.append(f"{summary}")

        page_content = "\n".join(doc_parts)

        document = Document(
            page_content=page_content,
            metadata={"source": source}
        )

        documents.append(document)

    return documents


print(load_document("https://hackrx.blob.core.windows.net/assets/policy.pdf?sv=2023-01-03&st=2025-07-04T09%3A11%3A24Z&se=2027-07-05T09%3A11%3A00Z&sr=b&sp=r&sig=N4a9OU0w0QXO6AOIBiu4bpl7AXvEZogeT%2FjUHNO7HzQ%3D"))

# Example Usage:
# data = {
#     "text": "This is a long paragraph of text...",
#     "tables": [
#         [["c1", "c2"], ["Organ Donor's Hospitalisation Expenses", None], ["Row3 Col1", "Row3 Col2"]],
#         [["Another Table Row1Col1", "Row1Col2"], ["Row2Col1", "Row2Col2"]]
#     ],
#     "image_summaries": [{'page': 1, 'image_index': 1, 'ocr_summary': 'ARe FRAT\n\nNational Insurance'}]
# }
# data=process_file("https://hackrx.blob.core.windows.net/assets/policy.pdf?sv=2023-01-03&st=2025-07-04T09%3A11%3A24Z&se=2027-07-05T09%3A11%3A00Z&sr=b&sp=r&sig=N4a9OU0w0QXO6AOIBiu4bpl7AXvEZogeT%2FjUHNO7HzQ%3D")

# document_1 = dict_to_document(source="document")
# print(document_1.page_content)
# print(document_1.metadata)
